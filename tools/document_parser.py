"""统一文档解析层，支持 PDF、DOCX、TXT 和 MD。"""

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Tuple

from app.config import KNOWLEDGE_EXTENSIONS
from tools.text_quality import TextQualityValidator


@dataclass
class ParsedDocument:
    """统一文档解析结果。"""

    text: str
    file_type: str
    parser_name: str
    text_length: int
    quality_score: float
    warnings: List[str]
    metadata: Dict[str, object]


def parse_document(file_path: str) -> ParsedDocument:
    """兼容旧调用方式，使用统一文档解析服务返回解析结果。"""

    return DocumentParserService().parse(file_path)


class DocumentParserService:
    """统一文档解析服务，屏蔽 PDF、DOCX、TXT、MD 的格式差异。"""

    def __init__(self) -> None:
        """初始化解析服务使用的文本质量验证器。"""

        self.quality_validator = TextQualityValidator()

    def parse(self, file_path: str) -> ParsedDocument:
        """解析文件并返回标准化的 ParsedDocument 对象。"""

        path = Path(file_path)
        suffix = path.suffix.lower()
        if suffix not in KNOWLEDGE_EXTENSIONS:
            raise ValueError(f"不支持的文件类型: {suffix}")

        if suffix == ".pdf":
            raw_text, parser_name, metadata = _parse_pdf(path)
        elif suffix == ".docx":
            raw_text, parser_name, metadata = _parse_docx(path)
        elif suffix in {".txt", ".md"}:
            raw_text, parser_name, metadata = _parse_text(path)
        else:
            raise ValueError(f"无法解析文件类型: {suffix}")

        text = normalize_document_text(raw_text)
        quality = self.quality_validator.validate(text)
        return ParsedDocument(
            text=text,
            file_type=suffix.lstrip("."),
            parser_name=parser_name,
            text_length=int(quality["text_length"]),
            quality_score=float(quality["quality_score"]),
            warnings=list(quality["warnings"]),
            metadata={**metadata, "source": str(path)},
        )


def normalize_document_text(text: str) -> str:
    """规范化文档文本中的空格、制表符和连续空行。"""

    normalized = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"[ \t\u3000]+", " ", normalized)
    normalized = re.sub(r" *\n *", "\n", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def _parse_pdf(path: Path) -> tuple:
    """使用 pypdf 解析 PDF 文本。"""

    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("解析 PDF 需要安装 pypdf") from exc

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    metadata = {"page_count": len(reader.pages)}
    return "\n".join(pages), "pypdf", metadata


def _parse_docx(path: Path) -> tuple:
    """使用 python-docx 解析 DOCX 正文、表格、嵌套表格、页眉和页脚文本。"""

    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("解析 DOCX 需要安装 python-docx") from exc

    document = Document(str(path))
    parts: List[str] = []
    for paragraph in document.paragraphs:
        text = normalize_document_text(paragraph.text)
        if text:
            parts.append(text)

    header_footer_count = 0
    header_footer_table_count = 0
    header_footer_table_cell_count = 0
    for section in document.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                text = normalize_document_text(paragraph.text)
                if text:
                    parts.append(text)
                    header_footer_count += 1
            container_table_parts, container_cell_count = _extract_docx_table_text(container.tables)
            if container_table_parts:
                parts.extend(container_table_parts)
                header_footer_table_count += len(container.tables)
                header_footer_table_cell_count += container_cell_count

    table_parts, table_cell_count = _extract_docx_table_text(document.tables)
    parts.extend(table_parts)

    metadata = {
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "table_cell_count": table_cell_count,
        "header_footer_count": header_footer_count,
        "header_footer_table_count": header_footer_table_count,
        "header_footer_table_cell_count": header_footer_table_cell_count,
    }
    return "\n".join(parts), "python-docx-paragraphs-tables-headers-footers", metadata


def _extract_docx_table_text(tables: list) -> Tuple[List[str], int]:
    """提取 DOCX 表格、单元格段落和嵌套表格中的文本。"""

    parts: List[str] = []
    table_cell_count = 0
    for table in tables:
        for row in table.rows:
            row_values = []
            for cell in row.cells:
                cell_parts = _extract_docx_cell_text(cell)
                if cell_parts:
                    row_values.append(" ".join(cell_parts))
                    table_cell_count += 1
                nested_parts, nested_count = _extract_docx_table_text(cell.tables)
                parts.extend(nested_parts)
                table_cell_count += nested_count
            if row_values:
                parts.append(" | ".join(row_values))
    return parts, table_cell_count


def _extract_docx_cell_text(cell) -> List[str]:
    """提取单元格内普通段落文本，避免只读 cell.text 时丢失复杂排版内容。"""

    values: List[str] = []
    for paragraph in cell.paragraphs:
        text = normalize_document_text(paragraph.text)
        if text:
            values.append(text)
    if not values:
        fallback = normalize_document_text(cell.text)
        if fallback:
            values.append(fallback)
    return values


def _parse_text(path: Path) -> tuple:
    """解析 TXT 或 MD 文本文件。"""

    try:
        text = path.read_text(encoding="utf-8")
        encoding = "utf-8"
    except UnicodeDecodeError:
        text = path.read_text(encoding="gbk", errors="ignore")
        encoding = "gbk"
    return text, f"text-{encoding}", {"encoding": encoding}
