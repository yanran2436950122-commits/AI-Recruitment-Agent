"""不同文档格式解析与评分一致性测试。"""

from pathlib import Path

from tools.document_parser import parse_document
from tools.score_tool import calculate_match_score, extract_skills


DEMO_RESUME_TEXT = """
Candidate: Zhang San
Target Role: AI Agent Engineer
Skills: Python, FastAPI, LangGraph, RAG, Redis, PostgreSQL, Docker, Vector Database, Chroma, Milvus.
Project Experience: Built an enterprise recruitment agent platform with resume parsing, JD matching,
RAG retrieval, vector database integration, memory layer, FastAPI service APIs, and Streamlit workflow UI.
Project Result: Improved candidate screening efficiency, generated customized interview questions,
produced recruitment reports, and supported HR candidate comparison.
Education: Computer Science related major.
""".strip()

DEMO_JD_TEXT = (
    "Recruiting an AI Agent Engineer. Required skills include Python, FastAPI, LangGraph, "
    "RAG, Redis, PostgreSQL, Docker, vector database, Chroma, Milvus, resume parsing, "
    "JD matching, interview question generation, and recruitment report experience."
)


def test_txt_docx_pdf_parse_and_score_consistency(tmp_path: Path) -> None:
    """同内容 TXT、DOCX 表格和 PDF 简历应解析出接近文本，并得到接近基础分。"""

    txt_path = tmp_path / "demo_resume.txt"
    docx_path = tmp_path / "demo_resume.docx"
    pdf_path = tmp_path / "demo_resume.pdf"
    txt_path.write_text(DEMO_RESUME_TEXT, encoding="utf-8")
    create_demo_docx(docx_path)
    create_demo_pdf(pdf_path, DEMO_RESUME_TEXT)

    txt_doc = parse_document(str(txt_path))
    docx_doc = parse_document(str(docx_path))
    pdf_doc = parse_document(str(pdf_path))
    txt_score, _, _ = calculate_match_score(txt_doc.text, DEMO_JD_TEXT)
    docx_score, _, _ = calculate_match_score(docx_doc.text, DEMO_JD_TEXT)
    pdf_score, _, _ = calculate_match_score(pdf_doc.text, DEMO_JD_TEXT)

    assert docx_doc.text_length > 200
    assert pdf_doc.text_length > 200
    assert "python" in docx_doc.text.lower()
    assert "python" in pdf_doc.text.lower()
    assert "langgraph" in docx_doc.text.lower()
    assert "langgraph" in pdf_doc.text.lower()
    assert set(extract_skills(txt_doc.text)).intersection(extract_skills(docx_doc.text))
    assert set(extract_skills(txt_doc.text)).intersection(extract_skills(pdf_doc.text))
    assert abs(txt_doc.text_length - docx_doc.text_length) <= 120
    assert abs(txt_doc.text_length - pdf_doc.text_length) <= 120
    assert abs(txt_score - docx_score) <= 5
    assert abs(pdf_score - docx_score) <= 5


def create_demo_docx(path: Path) -> None:
    """创建包含表格内容的 DOCX 简历样例，用于验证单元格文本不会丢失。"""

    from docx import Document

    document = Document()
    document.add_heading("Candidate Resume", level=1)
    table = document.add_table(rows=6, cols=2)
    rows = [
        ("Candidate", "Zhang San"),
        ("Target Role", "AI Agent Engineer"),
        ("Skills", "Python, FastAPI, LangGraph, RAG, Redis, PostgreSQL, Docker, Vector Database, Chroma, Milvus"),
        (
            "Project Experience",
            "Built an enterprise recruitment agent platform with resume parsing, JD matching, "
            "RAG retrieval, vector database integration, memory layer, FastAPI service APIs, and Streamlit workflow UI.",
        ),
        (
            "Project Result",
            "Improved candidate screening efficiency, generated customized interview questions, "
            "produced recruitment reports, and supported HR candidate comparison.",
        ),
        ("Education", "Computer Science related major."),
    ]
    for row, values in zip(table.rows, rows):
        row.cells[0].text = values[0]
        row.cells[1].text = values[1]
    document.save(str(path))


def create_demo_pdf(path: Path, text: str) -> None:
    """使用 pypdf 构造极简文本 PDF，避免测试额外依赖第三方 PDF 生成库。"""

    from pypdf import PdfWriter
    from pypdf.generic import DictionaryObject, NameObject, StreamObject

    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
            NameObject("/Encoding"): NameObject("/WinAnsiEncoding"),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): writer._add_object(font)})}
    )
    stream = StreamObject()
    stream._data = build_pdf_text_stream(text)
    page[NameObject("/Contents")] = writer._add_object(stream)
    with path.open("wb") as file:
        writer.write(file)


def build_pdf_text_stream(text: str) -> bytes:
    """把多行英文简历文本转换为 PDF 内容流，便于 pypdf 后续提取。"""

    commands = ["BT", "/F1 11 Tf", "72 740 Td"]
    for index, line in enumerate(text.splitlines()):
        if index:
            commands.append("0 -16 Td")
        commands.append(f"({escape_pdf_text(line)}) Tj")
    commands.append("ET")
    return "\n".join(commands).encode("latin-1")


def escape_pdf_text(text: str) -> str:
    """转义 PDF 字符串中的括号和反斜杠。"""

    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
