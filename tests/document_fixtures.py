"""文档解析与工作流测试使用的简历样例生成工具。"""

from pathlib import Path


def write_demo_docx(path: Path, text: str) -> None:
    """把多行简历文本写入带表格的 DOCX 文件。"""

    from docx import Document

    document = Document()
    document.add_heading("Resume", level=1)
    table = document.add_table(rows=len(text.splitlines()), cols=2)
    for row, line in zip(table.rows, text.splitlines()):
        label, _, value = line.partition(":")
        row.cells[0].text = label.strip() or "Content"
        row.cells[1].text = value.strip() or line.strip()
    document.sections[0].header.paragraphs[0].text = "Resume Header Skill Summary"
    document.sections[0].footer.paragraphs[0].text = "Resume Footer Project Summary"
    document.save(str(path))


def write_demo_pdf(path: Path, text: str) -> None:
    """使用 pypdf 构造包含可提取文本的极简 PDF 文件。"""

    from pypdf import PdfWriter
    from pypdf.generic import DictionaryObject, NameObject, StreamObject

    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
            NameObject("/Encoding"): NameObject("/WinAnsiEncoding"),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): writer._add_object(font)})}
    )
    stream = StreamObject()
    stream._data = _build_pdf_text_stream(text)
    page[NameObject("/Contents")] = writer._add_object(stream)
    with path.open("wb") as file:
        writer.write(file)


def _build_pdf_text_stream(text: str) -> bytes:
    """把多行英文文本转换为 PDF 内容流。"""

    commands = ["BT", "/F1 11 Tf", "72 740 Td"]
    for index, line in enumerate(text.splitlines()):
        if index:
            commands.append("0 -16 Td")
        commands.append(f"({_escape_pdf_text(line)}) Tj")
    commands.append("ET")
    return "\n".join(commands).encode("latin-1")


def _escape_pdf_text(text: str) -> str:
    """转义 PDF 文本中的括号和反斜杠。"""

    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
