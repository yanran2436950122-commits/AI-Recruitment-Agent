"""DOCX 解析覆盖率与文本质量评分测试。"""

from pathlib import Path

from tools.document_parser import parse_document
from tools.text_quality import validate_resume_text


def test_docx_parser_extracts_tables_headers_footers_and_nested_tables(tmp_path: Path) -> None:
    """DOCX 解析应覆盖正文表格、嵌套表格、页眉和页脚内容。"""

    docx_path = tmp_path / "complex_resume.docx"
    create_complex_docx(docx_path)

    parsed = parse_document(str(docx_path))
    text = parsed.text.lower()

    assert "python" in text
    assert "langgraph" in text
    assert "nested rag project" in text
    assert "header skill fastapi" in text
    assert "footer project summary" in text
    assert parsed.metadata["table_cell_count"] >= 2
    assert parsed.metadata["header_footer_table_cell_count"] >= 1
    assert parsed.quality_score >= 0.3


def test_low_quality_resume_text_returns_warnings() -> None:
    """过短且缺少技能/项目经历的文本应返回质量 warning。"""

    quality = validate_resume_text("姓名：张三")

    assert quality["text_length"] < 200
    assert quality["quality_score"] < 0.7
    assert any("长度小于 200" in warning for warning in quality["warnings"])
    assert any("技能关键词过少" in warning for warning in quality["warnings"])
    assert any("项目经历关键词过少" in warning for warning in quality["warnings"])


def create_complex_docx(path: Path) -> None:
    """创建包含复杂排版元素的 DOCX 简历。"""

    from docx import Document

    document = Document()
    document.add_paragraph("Candidate: Complex Docx User")
    document.sections[0].header.paragraphs[0].text = "Header Skill FastAPI"
    document.sections[0].footer.paragraphs[0].text = "Footer Project Summary"
    header_table = document.sections[0].header.add_table(rows=1, cols=1, width=1000000)
    header_table.cell(0, 0).text = "Header Table Skill Redis"

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "Python, LangGraph, RAG, Docker, PostgreSQL, Milvus"
    table.cell(1, 0).text = "Project Experience"
    nested_table = table.cell(1, 1).add_table(rows=1, cols=1)
    nested_table.cell(0, 0).text = "Nested RAG Project with vector database retrieval and interview generation"
    document.save(str(path))
